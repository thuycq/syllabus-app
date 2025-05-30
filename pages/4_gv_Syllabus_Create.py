import requests
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date
import streamlit as st
import os
import pandas as pd


# Tải ảnh logo từ Google Drive
logo_url = "https://drive.google.com/uc?export=download&id=1TunxDkuBhjAjkNYwuYGEAkVXOyN2RJdO"
logo_path = "logo_temp.png"
with open(logo_path, "wb") as f:
    f.write(requests.get(logo_url).content)


st.set_page_config(page_title="Syllabus Management", page_icon="📘", layout="wide")
st.title("Soạn thảo đề cương môn họcc")


tabs = st.tabs(['1. Thông tin tổng quát', '2. Mô tả và Giáo trình', '3. Mục tiêu môn học', '4. Chuẩn đầu ra môn học', '5. Đánh giá môn học', '6. Kế hoạch giảng dạy chi tiết', '7. Thời lượng học tập', '8. Các nội dung khác'])


# --- 1. Thông tin tổng quát ---
with tabs[0]:  # 1. Thông tin tổng quát
    col1, col2 = st.columns(2)
    with col1:
        ten_tv = st.text_input("Tên học phần (tiếng Việt)")
        ma_mh = st.text_input("Mã học phần")
        nganh_hoc = st.text_input("Ngành học")
        so_tc = st.text_input("Số tín chỉ")
        so_tiet_lt = st.text_input("Số tiết lý thuyết", value="45")
        so_tiet_th = st.text_input("Số tiết thực hành", value="0")
        so_gio_tuhoc = st.text_input("Số giờ tự học", value="135")
        mon_hoc_truoc = st.text_input("Môn học trước", value="Không")
        mon_hoc_sau = st.text_input("Môn học sau", value="Không")
        mon_hoc_tien_quyet = st.text_input("Môn học tiên quyết", value="Không")
        mon_hoc_song_hanh = st.text_input("Môn học song hành", value="Không")
    with col2:
        ten_ta = st.text_input("Tên học phần (tiếng Anh)")
        trinh_do = st.selectbox(
            "Trình độ", 
            ["Đại học", "Thạc sĩ"],
        )
        thuoc_kien_thuc = st.selectbox("Thuộc kiến thức", [
            "Kiến thức đại cương", "Kiến thức cơ sở ngành", "Kiến thức ngành", "Kiến thức chuyên ngành"
        ])
        loai_mon = st.selectbox("Loại môn", ["Bắt buộc", "Tự chọn"])
        ngon_ngu = st.text_input("Ngôn ngữ giảng dạy", "Tiếng Việt")
        giang_vien = st.text_input("Giảng viên phụ trách")
        tro_giang = st.text_input("Giảng viên trợ giảng")
        nam_hoc = st.text_input("Năm học")
        hoc_ky = st.text_input("Học kỳ")
        if trinh_do == "Thạc sĩ":
            khoa_hoc = st.selectbox(
                "Khóa học", 
                ["25"],
            )
            ctdt_options = [
                "Chuyên ngành Tài chính - Ngân hàng định hướng Nghiên cứu",
                "Chuyên ngành Tài chính - Ngân hàng định hướng Ứng dụng",
                "Chuyên ngành Công nghệ tài chính định hướng Nghiên cứu",
                "Chuyên ngành Công nghệ tài chính định hướng Ứng dụng"
            ]
        elif trinh_do == "Đại học":
            khoa_hoc = st.selectbox(
                "Khóa học", 
                ["23", "25"], 
            )
            if khoa_hoc == "23":
                ctdt_options = [
                    "Tài chính - Ngân hàng",
                    "Tài chính - Ngân hàng Tiếng Anh",
                    "Công nghệ tài chính"
                ]
            elif khoa_hoc == "25":
                ctdt_options = [
                    "Tài chính - Ngân hàng",
                    "Tài chính - Ngân hàng Tiếng Anh",
                    "Công nghệ tài chính",
                    "Công nghệ tài chính liên kết doanh nghiệp"
                ]
            else:
                ctdt_options = []

        ctdt = st.selectbox(
            "Chương trình đào tạo (CTĐT)", 
            ctdt_options,
        )


# --- 2+3. Mô tả & Tài liệu ---
with tabs[1]:  # 2. Mô tả môn học & Tài liệu học tập
    mo_ta = st.text_area("Mô tả môn học", value="(Nhập mô tả môn học)", height=150)
    #st.subheader("Tài liệu học tập")
    col1, col2 = st.columns(2)
    with col1:
        giao_trinh = st.text_area("Giáo trình", value="(Nhập ít nhất 2 Giáo trình chính)", height=100)
    with col2:
        tailieu_khac = st.text_area("Tài liệu khác", value="(Nhập các tài liệu tham khảo)", height=100)

#--- 4. Mục tiêu môn học ----
with tabs[2]:  # 4. Mục tiêu môn học
    muc_tieu_dict = st.data_editor(
        {"Mục tiêu (COx)": [""], "Mô tả mục tiêu": [""], "CĐR của MH (CLOx)": [""], "TĐNL": [""]},
        num_rows="dynamic"
    )
    muc_tieu_data = pd.DataFrame(muc_tieu_dict)

# --- 5. Chuẩn đầu ra môn học ---
with tabs[3]:  # 5. Chuẩn đầu ra môn học

    chuan_dau_ra_dict = st.data_editor(
        {
            "CĐR": [""],
            "Mô tả CĐR": [""],
            "Chuẩn đầu ra CTĐT (PLOx - PI x.x)": [""],
            "Thang Bloom/ Mức độ giảng dạy": [""]
        },
        num_rows="dynamic"
    )
    chuan_dau_ra_data = pd.DataFrame(chuan_dau_ra_dict)

# --- 6. Đánh giá kết quả học tập ---
with tabs[4]:  # 6. Đánh giá kết quả học tập
    # Tiểu mục 1
    yeu_cau_chung_mac_dinh = """+ Sinh viên vắng mặt trong buổi thi, đánh giá không có lý do chính đáng phải nhận điểm 0. Sinh viên vắng mặt có lý do chính đáng được dự thi, đánh giá ở một đợt khác và được tính điểm lần đầu.
+ Điểm thành phần: kết quả học tập một học phần được đánh giá theo quá trình học tập, thể hiện bởi các điểm thành phần được làm tròn tới một chữ số thập phân.
+ Điểm học phần được tính từ tổng các điểm thành phần nhân với trọng số tương ứng, được làm tròn tới một chữ số thập phân."""
    yeu_cau_chung = st.text_area(
        "Yêu cầu chung của học phần theo quy chế",
        value=yeu_cau_chung_mac_dinh,
        height=120
    )

    # Tiểu mục 2: Bảng đánh giá
    st.markdown("**Bảng đánh giá thành phần:**")
    bang_danh_gia_dict = st.data_editor(
        {
            "Thành phần đánh giá": [""],
            "Bài đánh giá": [""],
            "CĐR môn học": [""],
            "Tiêu chí đánh giá": [""],
            "Thời lượng đánh giá": [""],
            "Trọng số": [""],
            "Trọng số con": [""]
        },
        num_rows="dynamic"
    )
    bang_danh_gia_data = pd.DataFrame(bang_danh_gia_dict)

    # thang điểm và ghi chú
    ghi_chu_thang_diem = st.text_input("Ghi chú", value="- Học phần có điểm từ 5,0 trở lên được xem là học phần đạt, số tín chỉ của học phần này được tính là số tín chỉ tích lũy.")

# --- 7. Kế hoạch giảng dạy chi tiết ---
with tabs[5]:  # 7. Kế hoạch giảng dạy chi tiết
    st.markdown("### 📘 Lý thuyết")
    ke_hoach_lt_dict = st.data_editor(
        {
            "Tuần/Buổi": [""],
            "Nội dung": [""],
            "CĐR môn học": [""],
            "Hoạt động dạy và học": [""],
            "Hoạt động đánh giá": [""]
        },
        num_rows="dynamic",
        key="ke_hoach_lt"
    )
    ke_hoach_lt_data = pd.DataFrame(ke_hoach_lt_dict)

    st.markdown("### 🧪 Thực hành")
    ke_hoach_th_dict = st.data_editor(
        {
            "Tuần/Buổi": [""],
            "Nội dung": [""],
            "CĐR môn học": [""],
            "Hoạt động dạy và học": [""],
            "Hoạt động đánh giá": [""]
        },
        num_rows="dynamic",
        key="ke_hoach_th"
    )
    ke_hoach_th_data = pd.DataFrame(ke_hoach_th_dict)

# --- 8. Tổng thời lượng học tập ---
with tabs[6]:  # 8. Tổng thời lượng học tập
    #st.markdown("💡 **Điền thông tin theo từng dòng hoạt động**")

    tong_thoi_luong_dict = st.data_editor(
        {
            "Hình thức": [
                "Thời lượng học trên lớp (bao gồm tuần thi)",
                "Tự học ngoài giờ",
                "Bài tập tình huống",
                "Bài luận",
                "Bài thuyết trình",
                "Dự án nhóm",
                "Khác (ghi rõ)",
                "Thi giữa kỳ",
                "Thi cuối kỳ",
                "", "", ""
            ],
            "Hoạt động dạy và học": [
                "", "", "", "", "", "", "", "", "",
                "Tổng thời lượng",
                "Tổng thời lượng/… giờ",
                "Chuyển đổi tín chỉ theo ECTS"
            ],
            "Số lần": [""] * 12,
            "Thời lượng (giờ)": [""] * 12,
            "Tổng thời lượng (giờ)": [""] * 12
        },
        use_container_width=True,
        key="tong_thoi_luong",
        num_rows="fixed"
    )
    tong_thoi_luong_data = pd.DataFrame(tong_thoi_luong_dict)

# --- 9. Quy định và thông tin môn học (gom 9–12) ---
with tabs[7]:  # Quy định và thông tin môn học
    st.markdown("### • Quy định của môn học")
    quy_dinh_mac_dinh = """- Sinh viên nộp trễ hạn bài tập trên LMS sẽ bị trừ 50% số điểm của bài tập đó.
- Sinh viên vắng mặt trong buổi thi, đánh giá không có lý do chính đáng phải nhận điểm 0.
- Sinh viên vắng mặt có lý do chính đáng được dự thi, đánh giá ở một đợt khác và được tính điểm lần đầu."""
    quy_dinh_mh = st.text_area("Nội dung quy định", value=quy_dinh_mac_dinh, height=150)

    st.markdown("### Phụ trách môn học")
    col1, col2 = st.columns(2)
    with col1:
        khoa_phu_trach = st.text_input("Khoa", value="Tài chính - Ngân hàng")
        bo_mon_phu_trach = st.text_input("Bộ môn", value="")
    with col2:
        email_lien_he = st.text_input("Email liên hệ", value="")

    st.markdown("### • Ngày biên soạn và thẩm định")
    col3, col4 = st.columns(2)
    with col3:
        ngay_bien_soan = st.date_input("Ngày biên soạn/ cập nhật")
    with col4:
        ngay_tham_dinh = st.date_input("Ngày thẩm định/ thông qua")


# --- Export Function ---
os.makedirs("syllabus", exist_ok=True)

def export_syllabus_to_word(
    muc_tieu_data,
    chuan_dau_ra_data,
    bang_danh_gia_data,
    yeu_cau_chung,
    ke_hoach_lt_data,
    ke_hoach_th_data,
    tong_thoi_luong_data,
    quy_dinh_mh,
    khoa_phu_trach,
    bo_mon_phu_trach,
    email_lien_he,
    ngay_bien_soan,
    ngay_tham_dinh,
    ma_mh,
    trinh_do,
    ten_tv,
    khoa_hoc
):
    doc = Document()

    # Font mặc định: Times New Roman, size 13
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    rPr = style._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    rPr.append(rFonts)

    # Logo + Tiêu ngữ
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header_table.autofit = False

    logo_cell = header_table.cell(0, 0)
    logo_cell.width = Inches(1.2)
    logo_paragraph = logo_cell.paragraphs[0]
    logo_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    logo_paragraph.add_run().add_picture(logo_path, width=Inches(1.1))

    text_cell = header_table.cell(0, 1)
    text_cell.width = Inches(3.2)
    for line in ["TRƯỜNG ĐẠI HỌC KINH TẾ - LUẬT",
        "KHOA TÀI CHÍNH NGÂN HÀNG",
        "BỘ MÔN NGÂN HÀNG"
    ]:
        p = text_cell.add_paragraph(line)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        r = p.runs[0]
        if "KHOA" in line or "BỘ MÔN" in line:
            r.bold = True

    doc.add_paragraph("")

    # Tiêu đề chính
    p = doc.add_paragraph("Đề cương chi tiết học phần".title())
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.runs[0].bold = True

    doc.add_paragraph("")

    # 1. Thông tin tổng quát
    doc.add_paragraph("1. Thông tin tổng quát:").runs[0].bold = True

    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    def add_row(label, value):
        row = table.add_row().cells
        run = row[0].paragraphs[0].add_run(label)
        run.bold = False
        row[1].text = str(value)

    add_row("Tên học phần tiếng Việt", ten_tv)
    add_row("Tên học phần tiếng Anh", ten_ta)
    add_row("Mã học phần", ma_mh)
    add_row("Thuộc khối kiến thức", thuoc_kien_thuc)
    add_row("Loại môn học", loai_mon)
    add_row("Trình độ", trinh_do)
    add_row("Ngành học", nganh_hoc)
    add_row("Khóa học", khoa_hoc)
    add_row("Năm học", nam_hoc)
    add_row("Học kỳ", hoc_ky)
    add_row("Số tín chỉ", f"{so_tc} tín chỉ\nLý thuyết: {so_tiet_lt} tiết\nThực hành: {so_tiet_th} tiết\nTự học: {so_gio_tuhoc} giờ")
    add_row("Môn học tiên quyết", mon_hoc_tien_quyet)
    add_row("Môn học trước", mon_hoc_truoc)
    add_row("Môn học sau", mon_hoc_sau)
    add_row("Môn học song hành", mon_hoc_song_hanh)
    add_row("Ngôn ngữ sử dụng trong giảng dạy", ngon_ngu)
    add_row("Các giảng viên phụ trách giảng dạy", giang_vien)
    add_row("Các giảng viên trợ giảng", tro_giang)
   
    # --- 2. Mô tả môn học ---
    p2 = doc.add_paragraph("2. Mô tả môn học:")
    p2.runs[0].bold = True
    doc.add_paragraph(mo_ta)

    # --- 3. Tài liệu học tập ---
    p3 = doc.add_paragraph("3. Tài liệu học tập:")
    p3.runs[0].bold = True
    doc.add_paragraph("Giáo trình:")
    doc.add_paragraph(giao_trinh)
    doc.add_paragraph("Tài liệu khác:")
    doc.add_paragraph(tailieu_khac)

    # --- 4. Mục tiêu môn học ---
    p4 = doc.add_paragraph("4. Mục tiêu môn học:")
    p4.runs[0].bold = True

    mt_table = doc.add_table(rows=1, cols=4)
    mt_table.style = "Table Grid"
    hdr_cells = mt_table.rows[0].cells
    hdr_cells[0].text = "Mục tiêu (COx)"
    hdr_cells[1].text = "Mô tả mục tiêu"
    hdr_cells[2].text = "CĐR của MH (CLOx)"
    hdr_cells[3].text = "TĐNL"

    for _, row in muc_tieu_data.iterrows():
        cells = mt_table.add_row().cells
        cells[0].text = str(row["Mục tiêu (COx)"])
        cells[1].text = str(row["Mô tả mục tiêu"])
        cells[2].text = str(row["CĐR của MH (CLOx)"])
        cells[3].text = str(row["TĐNL"])

    # --- 5. Chuẩn đầu ra môn học ---
    p5 = doc.add_paragraph("5. Chuẩn đầu ra môn học:")
    p5.runs[0].bold = True
    doc.add_paragraph("Sau khi hoàn thành môn học, Người học có thể:")

    clo_table = doc.add_table(rows=1, cols=4)
    clo_table.style = "Table Grid"
    hdr = clo_table.rows[0].cells
    hdr[0].text = "CĐR (CLOx)"
    hdr[1].text = "Mô tả CĐR"
    hdr[2].text = "Chuẩn đầu ra CTĐT (PLOx)"
    hdr[3].text = "Mức độ đạt (I, R, M)"

    for _, row in chuan_dau_ra_data.iterrows():
        cells = clo_table.add_row().cells
        cells[0].text = str(row["CĐR"])
        cells[1].text = str(row["Mô tả CĐR"])
        cells[2].text = str(row["Chuẩn đầu ra CTĐT (PLOx - PI x.x)"])
        cells[3].text = str(row["Thang Bloom/ Mức độ giảng dạy"])
    # --- 6. Đánh giá kết quả học tập ---
    p6 = doc.add_paragraph("6. Đánh giá kết quả học tập:")
    p6.runs[0].bold = True

    # Tiểu mục 1
    doc.add_paragraph("- Yêu cầu chung của học phần theo quy chế:")
    doc.add_paragraph(yeu_cau_chung)

    # Tiểu mục 2
    doc.add_paragraph("Bảng đánh giá thành phần:")
    dg_table = doc.add_table(rows=1, cols=7)
    dg_table.style = "Table Grid"
    hdr = dg_table.rows[0].cells
    hdr[0].text = "Thành phần đánh giá"
    hdr[1].text = "Bài đánh giá"
    hdr[2].text = "CĐR môn học"
    hdr[3].text = "Tiêu chí đánh giá"
    hdr[4].text = "Thời lượng đánh giá"
    hdr[5].text = "Trọng số"
    hdr[6].text = "Trọng số con"

    for _, row in bang_danh_gia_data.iterrows():
        cells = dg_table.add_row().cells
        cells[0].text = str(row["Thành phần đánh giá"])
        cells[1].text = str(row["Bài đánh giá"])
        cells[2].text = str(row["CĐR môn học"])
        cells[3].text = str(row["Tiêu chí đánh giá"])
        cells[4].text = str(row["Thời lượng đánh giá"])
        cells[5].text = str(row["Trọng số"])
        cells[6].text = str(row["Trọng số con"])

    # Tiểu mục 3
    doc.add_paragraph("- Thang điểm đánh giá: theo thang điểm 10, thang điểm 100 hoặc thang điểm 4 được dùng để ghi nhận và phân loại kết quả đánh giá. Thang điểm và cách xếp loại kết quả học tập được thực hiện như sau:")

    # Tạo bảng thang điểm
    grading_table = doc.add_table(rows=1, cols=5)
    grading_table.style = "Table Grid"
    hdr = grading_table.rows[0].cells
    hdr[0].text = "Xếp loại"
    hdr[1].text = "Thang điểm hệ 10"
    hdr[2].text = "Thang điểm hệ 100"
    hdr[3].text = "Thang điểm hệ 4\n(Điểm số)"
    hdr[4].text = "Thang điểm hệ 4\n(Điểm chữ)"

    # Dữ liệu bảng
    grading_data = [
        ["Xuất sắc", "Từ 9,0 đến 10,0", "Từ 90 đến 100", "4,0", "A+"],
        ["Giỏi", "Từ 8,0 đến cận 9,0", "Từ 80 đến cận 90", "3,5", "A"],
        ["Khá", "Từ 7,0 đến cận 8,0", "Từ 70 đến cận 80", "3,0", "B+"],
        ["Trung bình khá", "Từ 6,0 đến cận 7,0", "Từ 60 đến cận 70", "2,5", "B"],
        ["Trung bình", "Từ 5,0 đến cận 6,0", "Từ 50 đến cận 60", "2,0", "C"],
        ["Yếu", "Từ 4,0 đến cận 5,0", "Từ 40 đến cận 50", "1,5", "D+"],
        ["Kém", "Từ 3,0 đến cận 4,0", "Từ 30 đến cận 40", "1,0", "D"],
        ["", "< 3,0", "Dưới 30", "0,0", "F"],
    ]

    for row in grading_data:
        cells = grading_table.add_row().cells
        for i in range(5):
            cells[i].text = row[i]

    # Ghi chú bên dưới bảng
    doc.add_paragraph(ghi_chu_thang_diem,  style="Normal")

    # --- 7. Kế hoạch giảng dạy chi tiết ---
    doc.add_paragraph("7. Kế hoạch giảng dạy chi tiết:").runs[0].bold = True

    # Lý thuyết
    doc.add_paragraph("Lý thuyết:")
    lt_table = doc.add_table(rows=1, cols=5)
    lt_table.style = "Table Grid"
    hdr = lt_table.rows[0].cells
    hdr[0].text = "Tuần/Buổi"
    hdr[1].text = "Nội dung"
    hdr[2].text = "CĐR môn học"
    hdr[3].text = "Hoạt động dạy và học"
    hdr[4].text = "Hoạt động đánh giá"

    for _, row in ke_hoach_lt_data.iterrows():
        cells = lt_table.add_row().cells
        cells[0].text = str(row["Tuần/Buổi"])
        cells[1].text = str(row["Nội dung"])
        cells[2].text = str(row["CĐR môn học"])
        cells[3].text = str(row["Hoạt động dạy và học"])
        cells[4].text = str(row["Hoạt động đánh giá"])

    # Thực hành
    doc.add_paragraph("Thực hành:")
    th_table = doc.add_table(rows=1, cols=5)
    th_table.style = "Table Grid"
    hdr = th_table.rows[0].cells
    hdr[0].text = "Tuần/Buổi"
    hdr[1].text = "Nội dung"
    hdr[2].text = "CĐR môn học"
    hdr[3].text = "Hoạt động dạy và học"
    hdr[4].text = "Hoạt động đánh giá"

    for _, row in ke_hoach_th_data.iterrows():
        cells = th_table.add_row().cells
        cells[0].text = str(row["Tuần/Buổi"])
        cells[1].text = str(row["Nội dung"])
        cells[2].text = str(row["CĐR môn học"])
        cells[3].text = str(row["Hoạt động dạy và học"])
        cells[4].text = str(row["Hoạt động đánh giá"])

    # --- 8. Tổng thời lượng học tập ---
    doc.add_paragraph("8. Tổng thời lượng học tập:").runs[0].bold = True

    # Tạo bảng tiêu đề
    tong_table = doc.add_table(rows=1, cols=5)
    tong_table.style = "Table Grid"

    # Tiêu đề cột
    header_cells = tong_table.rows[0].cells
    header_cells[0].text = "Hình thức"
    header_cells[1].text = "Hoạt động dạy và học"
    header_cells[2].text = "Số lần"
    header_cells[3].text = "Thời lượng (giờ)"
    header_cells[4].text = "Tổng thời lượng (giờ)"

    # Ghi từng dòng dữ liệu
    for _, row in tong_thoi_luong_data.iterrows():
        row_cells = tong_table.add_row().cells
        row_cells[0].text = str(row.get("Hình thức", ""))
        row_cells[1].text = str(row.get("Hoạt động dạy và học", ""))
        row_cells[2].text = str(row.get("Số lần", ""))
        row_cells[3].text = str(row.get("Thời lượng (giờ)", ""))
        row_cells[4].text = str(row.get("Tổng thời lượng (giờ)", ""))

    # Ghi chú in nghiêng
    doc.add_paragraph().add_run(
        "Ghi chú: 1 tiết = 50 phút = 5/6 giờ; 1 tín chỉ 50h học tập bao gồm cả thời gian học tập trên lớp, tự học, nghiên cứu, dự kiểm tra, đánh giá."
    ).italic = True

    # --- 9. Quy định của môn học ---
    doc.add_paragraph("9. Quy định của môn học:").runs[0].bold = True
    doc.add_paragraph(quy_dinh_mh)

    # --- 10. Phụ trách môn học ---
    doc.add_paragraph("10. Phụ trách môn học:").runs[0].bold = True
    doc.add_paragraph(f"- Khoa: {khoa_phu_trach}")
    doc.add_paragraph(f"- Bộ môn: {bo_mon_phu_trach}")
    doc.add_paragraph(f"- Email liên hệ: {email_lien_he}")

    # --- 11. Ngày biên soạn & 12. Thẩm định ---
    para_11 = doc.add_paragraph()
    run_11 = para_11.add_run(f"11. Đề cương được cập nhật và biên soạn ngày: {ngay_bien_soan.strftime('%d/%m/%Y')}")
    run_11.bold = True

    para_12 = doc.add_paragraph()
    run_12 = para_12.add_run(f"12. Đề cương được thẩm định và thông qua ngày: {ngay_tham_dinh.strftime('%d/%m/%Y')}")
    run_12.bold = True

    # --- Ký tên ---
    sign_table = doc.add_table(rows=2, cols=2)
    sign_table.columns[0].width = Inches(3)
    sign_table.columns[1].width = Inches(3.5)

    # Hàng 1 - thời gian
    sign_table.cell(0, 0).width = Inches(3)
    sign_table.cell(0, 1).width = Inches(4.5)
    sign_table.cell(0, 0).text = ""
    p_time = sign_table.cell(0, 1).paragraphs[0]
    p_time.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_time = p_time.add_run("TP. Hồ Chí Minh, ngày …… tháng …… năm ……")
    run_time.font.size = Pt(11)
    run_time.italic = True

    # Hàng 2 - ký tên
    sign_table.cell(1, 0).width = Inches(3)
    sign_table.cell(1, 1).width = Inches(4.5)
    p1 = sign_table.cell(1, 0).paragraphs[0]
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p1.add_run("Giảng viên biên soạn")

    p2 = sign_table.cell(1, 1).paragraphs[0]
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p2.add_run("Trưởng khoa/Bộ môn")

   
    # Xác định folder lưu
    trinh_do_folder = ""
    if trinh_do.lower() == "đại học":
        trinh_do_folder = "daihoc"
    elif trinh_do.lower() == "thạc sĩ":
        trinh_do_folder = "thacsi"

    ctdt_folder = ctdt.strip().lower()

    folder_path = os.path.join("syllabus", trinh_do_folder, f"khoa{khoa_hoc}", ctdt_folder)
    os.makedirs(folder_path, exist_ok=True)

    file_name = f"{ma_mh}_ĐCCT_{ten_tv}_{khoa_hoc}.docx"
    full_file_path = os.path.join(folder_path, file_name)

    doc.save(full_file_path)

    return full_file_path

# --- Export Button ---
col1, col2 = st.columns(2)

with col1:
    if st.button("📄 Lưu và Xuất file"):
        if not ma_mh or not ten_tv or not khoa_hoc:
            st.error("⚠️ Bạn phải nhập đầy đủ Mã HP, Tên HP và Khóa trước khi lưu.")
        else:
            file_path = export_syllabus_to_word(
                muc_tieu_data,
                chuan_dau_ra_data,
                bang_danh_gia_data,
                yeu_cau_chung,
                ke_hoach_lt_data,
                ke_hoach_th_data,
                tong_thoi_luong_data,
                quy_dinh_mh,
                khoa_phu_trach,
                bo_mon_phu_trach,
                email_lien_he,
                ngay_bien_soan,
                ngay_tham_dinh,
                ma_mh,
                trinh_do,
                ten_tv,
                khoa_hoc
            )
            #st.success(f"✅ Đã lưu đề cương: {os.path.basename(file_path)}")

            # 👉 Thêm nút Tải xuống:
            with open(file_path, "rb") as f:
                st.download_button(
                    label="📥 Tải xuống Đề cương",
                    data=f.read(),
                    file_name=os.path.basename(file_path),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

# Nút Lưu và quay lại
with col2:
    if st.button("💾 Lưu và quay lại"):
        if not ma_mh or not ten_tv or not khoa_hoc:
            st.error("⚠️ Bạn phải nhập đầy đủ Mã HP, Tên HP và Khóa trước khi lưu.")
        else:
            file_path = export_syllabus_to_word(
                muc_tieu_data,
                chuan_dau_ra_data,
                bang_danh_gia_data,
                yeu_cau_chung,
                ke_hoach_lt_data,
                ke_hoach_th_data,
                tong_thoi_luong_data,
                quy_dinh_mh,
                khoa_phu_trach,
                bo_mon_phu_trach,
                email_lien_he,
                ngay_bien_soan,
                ngay_tham_dinh,
                ma_mh,
                trinh_do,
                ten_tv,
                khoa_hoc
            )
            st.success(f"✅ Đã lưu đề cương: {os.path.basename(file_path)}")
            st.switch_page("pages/2_gv_page.py")